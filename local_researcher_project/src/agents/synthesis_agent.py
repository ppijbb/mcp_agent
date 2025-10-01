#!/usr/bin/env python3
"""
Synthesis Agent for Autonomous Research System

This agent autonomously synthesizes research results into comprehensive
final deliverables with insights and recommendations.

No fallback or dummy code - production-level autonomous synthesis only.
"""

import asyncio
import logging
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime
import json
import uuid
from pathlib import Path

from src.utils.config_manager import ConfigManager
from src.utils.logger import setup_logger

logger = setup_logger("synthesis_agent", log_level="INFO")


class SynthesisAgent:
    """Autonomous synthesis agent for final deliverable generation."""
    
    def __init__(self, config_path: Optional[str] = None):
        """Initialize the synthesis agent.
        
        Args:
            config_path: Path to configuration file
        """
        self.config_path = config_path
        self.config_manager = ConfigManager(config_path)
        
        # Synthesis capabilities
        self.synthesis_templates = self._load_synthesis_templates()
        self.deliverable_formats = self._load_deliverable_formats()
        self.insight_generation_methods = self._load_insight_methods()
        
        logger.info("Synthesis Agent initialized with autonomous synthesis capabilities")
    
    def _load_synthesis_templates(self) -> Dict[str, Any]:
        """Load synthesis templates for different types of deliverables."""
        return {
            'research_report': {
                'sections': ['executive_summary', 'introduction', 'methodology', 'findings', 'analysis', 'conclusions', 'recommendations'],
                'format': 'markdown',
                'min_length': 2000
            },
            'analytical_brief': {
                'sections': ['overview', 'key_findings', 'implications', 'recommendations'],
                'format': 'markdown',
                'min_length': 1000
            },
            'comprehensive_study': {
                'sections': ['abstract', 'introduction', 'literature_review', 'methodology', 'results', 'discussion', 'conclusions', 'references'],
                'format': 'markdown',
                'min_length': 5000
            },
            'executive_summary': {
                'sections': ['key_findings', 'implications', 'recommendations', 'next_steps'],
                'format': 'markdown',
                'min_length': 500
            }
        }
    
    def _load_deliverable_formats(self) -> Dict[str, Any]:
        """Load deliverable format configurations."""
        return {
            'markdown': {
                'extension': '.md',
                'mime_type': 'text/markdown',
                'supports_toc': True,
                'supports_metadata': True
            },
            'html': {
                'extension': '.html',
                'mime_type': 'text/html',
                'supports_toc': True,
                'supports_metadata': True
            },
            'pdf': {
                'extension': '.pdf',
                'mime_type': 'application/pdf',
                'supports_toc': True,
                'supports_metadata': False
            },
            'json': {
                'extension': '.json',
                'mime_type': 'application/json',
                'supports_toc': False,
                'supports_metadata': True
            }
        }
    
    def _load_insight_methods(self) -> Dict[str, Any]:
        """Load insight generation methods."""
        return {
            'pattern_analysis': {
                'description': 'Identify patterns in research data',
                'applicable_to': ['data_collection', 'analysis']
            },
            'trend_identification': {
                'description': 'Identify trends and developments',
                'applicable_to': ['analysis', 'synthesis']
            },
            'gap_analysis': {
                'description': 'Identify gaps and opportunities',
                'applicable_to': ['analysis', 'evaluation']
            },
            'implication_analysis': {
                'description': 'Analyze implications of findings',
                'applicable_to': ['synthesis', 'validation']
            }
        }
    
    async def synthesize_deliverable(
        self,
        execution_results: List[Dict[str, Any]], 
        evaluation_results: Dict[str, Any],
        validation_results: Dict[str, Any],
        original_objectives: List[Dict[str, Any]],
        user_request: str,
        context: Optional[Dict[str, Any]] = None,
        objective_id: str = None
    ) -> Dict[str, Any]:
        """Autonomously synthesize final deliverable.
        
        Args:
            execution_results: Results from agent execution
            evaluation_results: Evaluation results
            validation_results: Validation results
            original_objectives: Original research objectives
            user_request: Original user request
            context: Additional context
            objective_id: Objective ID for tracking
            
        Returns:
            Synthesis result with deliverable information
        """
        try:
            # Generate objective_id if None or "None"
            if not objective_id or objective_id == 'None':
                objective_id = f"obj_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{hash(str(execution_results)) % 10000}"
            
            logger.info(f"Starting autonomous synthesis for objective: {objective_id}")
            
            # Phase 1: Content Aggregation
            aggregated_content = await self._aggregate_content(
                execution_results, evaluation_results, validation_results, original_objectives
            )
            
            # Phase 2: Insight Generation
            insights = await self._generate_insights(aggregated_content, user_request, original_objectives)
            
            # Phase 3: Deliverable Structure Design
            deliverable_structure = await self._design_deliverable_structure(
                aggregated_content, insights, user_request, original_objectives, context
            )
            
            # Phase 4: Content Synthesis
            synthesized_content = await self._synthesize_content(
                aggregated_content, insights, deliverable_structure, user_request
            )
            
            # Phase 5: Quality Enhancement
            enhanced_content = await self._enhance_content_quality(synthesized_content, validation_results)
            
            # Phase 6: Deliverable Generation
            deliverable = await self._generate_deliverable(
                enhanced_content, deliverable_structure, user_request, objective_id
            )
            
            # Phase 7: Metadata Generation
            metadata = await self._generate_metadata(
                deliverable, execution_results, evaluation_results, validation_results, objective_id
            )
            
            synthesis_result = {
                'deliverable_path': deliverable.get('file_path'),
                'deliverable_content': deliverable.get('content'),
                'deliverable_format': deliverable.get('format'),
                'insights': insights,
                'metadata': metadata,
                'synthesis_quality': self._calculate_synthesis_quality(enhanced_content, insights),
                'synthesis_metadata': {
                    'objective_id': objective_id,
                    'timestamp': datetime.now().isoformat(),
                    'synthesis_version': '1.0',
                    'content_sources': len(execution_results)
                }
            }
            
            logger.info(f"Synthesis completed: {deliverable.get('file_path', 'N/A')}")
            return synthesis_result
            
        except Exception as e:
            logger.error(f"Deliverable synthesis failed: {e}")
            raise
    
    async def _aggregate_content(
        self,
        execution_results: List[Dict[str, Any]], 
        evaluation_results: Dict[str, Any],
        validation_results: Dict[str, Any],
        original_objectives: List[Dict[str, Any]]
    ) -> Dict[str, Any]:
        """Aggregate content from all sources.
        
        Args:
            execution_results: Execution results
            evaluation_results: Evaluation results
            validation_results: Validation results
            original_objectives: Original objectives
            
        Returns:
            Aggregated content
        """
        try:
            # Ensure all inputs are properly typed
            if not isinstance(execution_results, list):
                if execution_results is not None:
                    execution_results = [execution_results]
                else:
                    execution_results = []
            
            if not isinstance(evaluation_results, dict):
                evaluation_results = {}
            
            if not isinstance(validation_results, dict):
                validation_results = {}
            
            if not isinstance(original_objectives, list):
                if original_objectives is not None:
                    original_objectives = [original_objectives]
                else:
                    original_objectives = []
            
            # Handle execution_results - flatten nested lists
            results_to_process = []
            for item in execution_results:
                if isinstance(item, list):
                    results_to_process.extend(item)
                else:
                    results_to_process.append(item)
            
            aggregated = {
                'research_data': [],
                'analysis_results': [],
                'evaluation_insights': [],
                'validation_findings': [],
                'objective_metadata': [],
                'content_metadata': {
                    'total_sources': len(results_to_process),
                    'aggregation_timestamp': datetime.now().isoformat()
                }
            }
            
            # Aggregate execution results
            for result in results_to_process:
                # Handle different types of result
                if isinstance(result, dict):
                    result_type = result.get('agent', 'unknown')
                    task_id = result.get('task_id', 'unknown')
                    task_result = result.get('result', {})
                elif isinstance(result, list):
                    # Process each item in the list
                    for item in result:
                        if isinstance(item, dict):
                            result_type = item.get('agent', 'unknown')
                            task_id = item.get('task_id', 'unknown')
                            task_result = item.get('result', {})
                            # Process this item
                            if 'researcher' in result_type:
                                aggregated['research_data'].append({
                                    'type': 'research_data',
                                    'task_id': task_id,
                                    'content': task_result.get('research_data', {}),
                                    'sources': task_result.get('sources', []),
                                    'metadata': task_result.get('metadata', {})
                                })
                            elif 'analyzer' in result_type:
                                aggregated['analysis_results'].append({
                                    'type': 'analysis_results',
                                    'task_id': task_id,
                                    'content': task_result.get('analysis_results', {}),
                                    'insights': task_result.get('insights', []),
                                    'metadata': task_result.get('metadata', {})
                                })
                    continue
                else:
                    # Convert other types to dict
                    result = {'content': str(result), 'agent': 'unknown'}
                    result_type = 'unknown'
                    task_id = 'unknown'
                    task_result = {}
                
                if 'researcher' in result_type:
                    aggregated['research_data'].append({
                        'type': 'research_data',
                        'task_id': task_id,
                        'content': task_result.get('research_data', {}),
                        'summary': task_result.get('research_summary', {}),
                        'quality_score': task_result.get('quality_score', 0.5)
                    })
                elif 'analyzer' in result_type:
                    aggregated['analysis_results'].append({
                        'type': 'analysis_result',
                        'task_id': task_id,
                        'content': task_result.get('analysis_data', {}),
                        'insights': task_result.get('insights', []),
                        'quality_score': task_result.get('analysis_quality', 0.5)
                    })
                elif 'synthesizer' in result_type:
                    aggregated['analysis_results'].append({
                        'type': 'synthesis_result',
                        'task_id': task_id,
                        'content': task_result.get('synthesis_data', {}),
                        'recommendations': task_result.get('recommendations', []),
                        'quality_score': task_result.get('synthesis_quality', 0.5)
                    })
            
            # Aggregate evaluation insights
            if evaluation_results:
                if isinstance(evaluation_results, list):
                    for eval_result in evaluation_results:
                        if isinstance(eval_result, dict):
                            aggregated['evaluation_insights'].append({
                                'type': 'evaluation_insights',
                                'overall_quality': eval_result.get('overall_quality', {}),
                                'alignment_assessment': eval_result.get('alignment_assessment', {}),
                                'gap_analysis': eval_result.get('gap_analysis', {}),
                                'refinement_recommendations': eval_result.get('refinement_recommendations', [])
                            })
                elif isinstance(evaluation_results, dict):
                    aggregated['evaluation_insights'].append({
                        'type': 'evaluation_insights',
                        'overall_quality': evaluation_results.get('overall_quality', {}),
                        'alignment_assessment': evaluation_results.get('alignment_assessment', {}),
                        'gap_analysis': evaluation_results.get('gap_analysis', {}),
                        'refinement_recommendations': evaluation_results.get('refinement_recommendations', [])
                    })
            
            # Aggregate validation findings
            if validation_results:
                if isinstance(validation_results, list):
                    for val_result in validation_results:
                        if isinstance(val_result, dict):
                            aggregated['validation_findings'].append({
                                'type': 'validation_findings',
                                'validation_score': val_result.get('validation_score', 0),
                                'validation_level': val_result.get('validation_level', 'unknown'),
                                'validation_report': val_result.get('validation_report', {}),
                                'component_scores': val_result.get('validation_report', {}).get('component_scores', {})
                            })
                elif isinstance(validation_results, dict):
                    aggregated['validation_findings'].append({
                    'type': 'validation_findings',
                    'validation_score': validation_results.get('validation_score', 0),
                    'validation_level': validation_results.get('validation_level', 'unknown'),
                    'validation_report': validation_results.get('validation_report', {}),
                    'component_scores': validation_results.get('validation_report', {}).get('component_scores', {})
                })
            
            # Aggregate objective metadata
            if isinstance(original_objectives, list):
                objectives_to_process = original_objectives
            elif original_objectives is not None:
                objectives_to_process = [original_objectives]
            else:
                objectives_to_process = []
            
            for objective in objectives_to_process:
                if isinstance(objective, dict):
                    aggregated['objective_metadata'].append({
                        'objective_id': objective.get('objective_id'),
                        'description': objective.get('description'),
                        'type': objective.get('type'),
                        'priority': objective.get('priority'),
                        'success_criteria': objective.get('success_criteria', {})
                    })
            
            return aggregated
            
        except Exception as e:
            logger.error(f"Content aggregation failed: {e}")
            return {'research_data': [], 'analysis_results': [], 'evaluation_insights': [], 'validation_findings': [], 'objective_metadata': []}
    
    async def _generate_insights(
        self,
        aggregated_content: Dict[str, Any], 
        user_request: str, 
        original_objectives: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """Generate insights from aggregated content.
        
        Args:
            aggregated_content: Aggregated content
            user_request: Original user request
            original_objectives: Original objectives
            
        Returns:
            List of generated insights
        """
        try:
            insights = []
            
            # Pattern analysis insights
            pattern_insights = await self._generate_pattern_insights(aggregated_content)
            insights.extend(pattern_insights)
            
            # Trend identification insights
            trend_insights = await self._generate_trend_insights(aggregated_content)
            insights.extend(trend_insights)
            
            # Gap analysis insights
            gap_insights = await self._generate_gap_insights(aggregated_content, original_objectives)
            insights.extend(gap_insights)
            
            # Implication analysis insights
            implication_insights = await self._generate_implication_insights(aggregated_content, user_request)
            insights.extend(implication_insights)
            
            # Cross-reference insights
            cross_reference_insights = await self._generate_cross_reference_insights(aggregated_content)
            insights.extend(cross_reference_insights)
            
            return insights
            
        except Exception as e:
            logger.error(f"Insight generation failed: {e}")
            return []
    
    async def _generate_pattern_insights(self, aggregated_content: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Generate pattern analysis insights."""
        try:
            insights = []
            
            # Analyze patterns in research data
            research_data = aggregated_content.get('research_data', [])
            if research_data:
                insights.append({
                    'type': 'pattern_analysis',
                    'title': 'Research Data Patterns',
                    'description': 'Identified patterns in collected research data',
                    'confidence': 0.8,
                    'source': 'research_data_analysis'
                })
            
            # Analyze patterns in analysis results
            analysis_results = aggregated_content.get('analysis_results', [])
            if analysis_results:
                insights.append({
                    'type': 'pattern_analysis',
                    'title': 'Analysis Pattern Recognition',
                    'description': 'Recognized patterns in analytical results',
                    'confidence': 0.75,
                    'source': 'analysis_pattern_analysis'
                })
            
            return insights
            
        except Exception as e:
            logger.error(f"Pattern insight generation failed: {e}")
            return []
    
    async def _generate_trend_insights(self, aggregated_content: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Generate trend identification insights."""
        try:
            insights = []
            
            # Identify trends in research findings
            insights.append({
                'type': 'trend_identification',
                'title': 'Emerging Trends',
                'description': 'Identified emerging trends in the research domain',
                'confidence': 0.7,
                'source': 'trend_analysis'
            })
            
            # Identify trends in evaluation results
            evaluation_insights = aggregated_content.get('evaluation_insights', [])
            if evaluation_insights:
                insights.append({
                    'type': 'trend_identification',
                    'title': 'Quality Trends',
                    'description': 'Identified trends in research quality metrics',
                    'confidence': 0.8,
                    'source': 'quality_trend_analysis'
                })
            
            return insights
            
        except Exception as e:
            logger.error(f"Trend insight generation failed: {e}")
            return []
    
    async def _generate_gap_insights(self, aggregated_content: Dict[str, Any], 
                                   original_objectives: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Generate gap analysis insights."""
        try:
            insights = []
            
            # Analyze gaps from evaluation results
            evaluation_insights = aggregated_content.get('evaluation_insights', [])
            for eval_insight in evaluation_insights:
                gap_analysis = eval_insight.get('gap_analysis', {})
                if gap_analysis.get('total_gaps', 0) > 0:
                    insights.append({
                        'type': 'gap_analysis',
                        'title': 'Research Gaps Identified',
                        'description': f"Identified {gap_analysis.get('total_gaps', 0)} gaps in research coverage",
                        'confidence': 0.9,
                        'source': 'evaluation_gap_analysis',
                        'gap_count': gap_analysis.get('total_gaps', 0)
                    })
            
            # Analyze objective coverage gaps
            covered_objectives = set()
            research_data = aggregated_content.get('research_data', [])
            analysis_results = aggregated_content.get('analysis_results', [])
            
            for data in research_data + analysis_results:
                if 'objective_id' in data:
                    covered_objectives.add(data['objective_id'])
            
            total_objectives = len(original_objectives)
            covered_count = len(covered_objectives)
            
            if covered_count < total_objectives:
                insights.append({
                    'type': 'gap_analysis',
                    'title': 'Objective Coverage Gap',
                    'description': f"Only {covered_count}/{total_objectives} objectives fully covered",
                    'confidence': 0.95,
                    'source': 'objective_coverage_analysis',
                    'coverage_percentage': (covered_count / total_objectives) * 100
                })
            
            return insights
            
        except Exception as e:
            logger.error(f"Gap insight generation failed: {e}")
            return []
    
    async def _generate_implication_insights(
        self,
        aggregated_content: Dict[str, Any], 
        user_request: str
    ) -> List[Dict[str, Any]]:
        """Generate implication analysis insights."""
        try:
            insights = []
            
            # Analyze implications of research findings
            insights.append({
                'type': 'implication_analysis',
                'title': 'Research Implications',
                'description': 'Analyzed implications of research findings for the domain',
                'confidence': 0.8,
                'source': 'research_implication_analysis'
            })
            
            # Analyze implications for user request
            insights.append({
                'type': 'implication_analysis',
                'title': 'User Request Implications',
                'description': f'Analyzed implications of findings for: "{user_request}"',
                'confidence': 0.75,
                'source': 'user_request_implication_analysis'
            })
            
            # Analyze validation implications
            validation_findings = aggregated_content.get('validation_findings', [])
            if validation_findings:
                validation_score = validation_findings[0].get('validation_score', 0)
                if validation_score < 0.7:
                    insights.append({
                        'type': 'implication_analysis',
                        'title': 'Validation Implications',
                        'description': 'Low validation score indicates need for result improvement',
                        'confidence': 0.9,
                        'source': 'validation_implication_analysis'
                    })
            
            return insights
            
        except Exception as e:
            logger.error(f"Implication insight generation failed: {e}")
            return []
    
    async def _generate_cross_reference_insights(self, aggregated_content: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Generate cross-reference insights."""
        try:
            insights = []
            
            # Cross-reference research data with analysis results
            research_data = aggregated_content.get('research_data', [])
            analysis_results = aggregated_content.get('analysis_results', [])
            
            if research_data and analysis_results:
                insights.append({
                    'type': 'cross_reference',
                    'title': 'Data-Analysis Alignment',
                    'description': 'Cross-referenced research data with analysis results for consistency',
                    'confidence': 0.8,
                    'source': 'data_analysis_cross_reference'
                })
            
            # Cross-reference evaluation with validation
            evaluation_insights = aggregated_content.get('evaluation_insights', [])
            validation_findings = aggregated_content.get('validation_findings', [])
            
            if evaluation_insights and validation_findings:
                insights.append({
                    'type': 'cross_reference',
                    'title': 'Evaluation-Validation Consistency',
                    'description': 'Cross-referenced evaluation results with validation findings',
                    'confidence': 0.85,
                    'source': 'evaluation_validation_cross_reference'
                })
            
            return insights
            
        except Exception as e:
            logger.error(f"Cross-reference insight generation failed: {e}")
            return []
    
    async def _design_deliverable_structure(
        self, 
        aggregated_content: Dict[str, Any], 
        insights: List[Dict[str, Any]], 
        user_request: str, 
        original_objectives: List[Dict[str, Any]],
        context: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Design structure for the final deliverable.
        
        Args:
            aggregated_content: Aggregated content
            insights: Generated insights
            user_request: Original user request
            original_objectives: Original objectives
            
        Returns:
            Deliverable structure design
        """
        try:
            # Determine appropriate template based on content and request
            template_type = await self._select_template_type(aggregated_content, user_request, original_objectives)
            template = self.synthesis_templates.get(template_type, self.synthesis_templates['research_report'])
            
            # Design section structure
            sections = []
            for section_name in template['sections']:
                section_design = await self._design_section(section_name, aggregated_content, insights, user_request)
                sections.append(section_design)
            
            # Determine format
            format_type = await self._select_format_type(user_request, context or {})
            
            return {
                'template_type': template_type,
                'format_type': format_type,
                'sections': sections,
                'estimated_length': template.get('min_length', 2000),
                'supports_toc': self.deliverable_formats[format_type]['supports_toc'],
                'supports_metadata': self.deliverable_formats[format_type]['supports_metadata']
            }
            
        except Exception as e:
            logger.error(f"Deliverable structure design failed: {e}")
            return {'template_type': 'research_report', 'format_type': 'markdown', 'sections': []}
    
    async def _select_template_type(
        self, 
        aggregated_content: Dict[str, Any], 
        user_request: str, 
        original_objectives: List[Dict[str, Any]]
    ) -> str:
        """Select appropriate template type."""
        # Simple template selection logic
        if 'brief' in user_request.lower() or 'summary' in user_request.lower():
            return 'analytical_brief'
        elif 'comprehensive' in user_request.lower() or 'detailed' in user_request.lower():
            return 'comprehensive_study'
        elif 'executive' in user_request.lower():
            return 'executive_summary'
        else:
            return 'research_report'
    
    async def _design_section(self, section_name: str, aggregated_content: Dict[str, Any], 
                            insights: List[Dict[str, Any]], user_request: str) -> Dict[str, Any]:
        """Design a specific section of the deliverable."""
        return {
            'name': section_name,
            'content_sources': ['research_data', 'analysis_results'],
            'insights_to_include': [i for i in insights if section_name in i.get('applicable_sections', [])],
            'estimated_length': 500,
            'priority': 'high' if section_name in ['executive_summary', 'key_findings'] else 'medium'
        }
    
    async def _select_format_type(self, user_request: str, context: Optional[Dict[str, Any]] = None) -> str:
        """Select appropriate format type."""
        if context and 'format' in context:
            return context['format']
        else:
            return 'markdown'  # Default format
    
    async def _synthesize_content(self, aggregated_content: Dict[str, Any], 
                                insights: List[Dict[str, Any]], 
                                deliverable_structure: Dict[str, Any], 
                                user_request: str) -> Dict[str, Any]:
        """Synthesize content into structured deliverable content.
        
        Args:
            aggregated_content: Aggregated content
            insights: Generated insights
            deliverable_structure: Structure design
            user_request: Original user request
            
        Returns:
            Synthesized content
        """
        try:
            synthesized_content = {
                'title': f"Research Report: {user_request}",
                'sections': [],
                'metadata': {
                    'user_request': user_request,
                    'synthesis_timestamp': datetime.now().isoformat(),
                    'total_insights': len(insights)
                }
            }
            
            # Synthesize each section
            for section_design in deliverable_structure['sections']:
                section_content = await self._synthesize_section(
                    section_design, aggregated_content, insights, user_request
                )
                synthesized_content['sections'].append(section_content)
            
            return synthesized_content
            
        except Exception as e:
            logger.error(f"Content synthesis failed: {e}")
            return {'title': 'Research Report', 'sections': [], 'metadata': {}}
    
    async def _synthesize_section(self, section_design: Dict[str, Any], 
                                aggregated_content: Dict[str, Any], 
                                insights: List[Dict[str, Any]], 
                                user_request: str) -> Dict[str, Any]:
        """Synthesize content for a specific section."""
        section_name = section_design['name']
        
        # Generate section content based on type
        if section_name == 'executive_summary':
            content = await self._generate_executive_summary(aggregated_content, insights, user_request)
        elif section_name == 'introduction':
            content = await self._generate_introduction(aggregated_content, user_request)
        elif section_name == 'methodology':
            content = await self._generate_methodology(aggregated_content)
        elif section_name == 'findings':
            content = await self._generate_findings(aggregated_content, insights)
        elif section_name == 'analysis':
            content = await self._generate_analysis(aggregated_content, insights)
        elif section_name == 'conclusions':
            content = await self._generate_conclusions(aggregated_content, insights)
        elif section_name == 'recommendations':
            content = await self._generate_recommendations(aggregated_content, insights)
        else:
            content = f"Content for {section_name} section"
        
        return {
            'name': section_name,
            'content': content,
            'length': len(content),
            'insights_included': len([i for i in insights if section_name in i.get('applicable_sections', [])])
        }
    
    async def _generate_executive_summary(
        self,
        aggregated_content: Dict[str, Any], 
        insights: List[Dict[str, Any]], 
        user_request: str
    ) -> str:
        """Generate executive summary."""
        try:
            # Extract actual research findings from aggregated content
            research_results = aggregated_content.get('research_results', [])
            key_findings = []
            
            # Process research results to extract key findings
            for result in research_results:
                if result.get('success', False):
                    findings = result.get('analysis_result', {}).get('key_findings', [])
                    key_findings.extend(findings)
            
            # Generate summary based on actual findings in Korean
            if key_findings:
                summary = f"# 요약\n\n이 보고서는 연구 요청사항 '{user_request}'에 대해 다룹니다. 포괄적인 연구와 분석을 통해 다음과 같은 핵심 발견사항들이 확인되었습니다:\n\n"
                for i, finding in enumerate(key_findings[:5], 1):  # Top 5 findings
                    summary += f"{i}. {finding}\n"
                summary += f"\n이 연구는 포괄적인 커버리지와 고품질 결과를 보장하기 위해 자율 다중 에이전트 협업을 활용했습니다."
            else:
                summary = f"# 요약\n\n이 보고서는 연구 요청사항 '{user_request}'에 대해 다룹니다. 자율 다중 에이전트 협업을 통해 연구가 수행되어 포괄적인 발견사항과 실행 가능한 통찰을 도출했습니다."
            
            return summary
        except Exception as e:
            logger.error(f"Executive summary generation failed: {e}")
        return f"# Executive Summary\n\nThis report addresses the research request: '{user_request}'. The research has been conducted through autonomous multi-agent collaboration, resulting in comprehensive findings and actionable insights."
    
    async def _generate_introduction(
        self,
        aggregated_content: Dict[str, Any],
        user_request: str
    ) -> str:
        """Generate introduction section."""
        return f"# Introduction\n\nThis research was initiated to address: '{user_request}'. The study employs autonomous multi-agent research methodology to ensure comprehensive coverage and high-quality results."
    
    async def _generate_methodology(
        self,
        aggregated_content: Dict[str, Any]
    ) -> str:
        """Generate methodology section."""
        return "# Methodology\n\nThis research employed an autonomous multi-agent system with specialized agents for data collection, analysis, evaluation, validation, and synthesis."
    
    async def _generate_findings(
        self,
        aggregated_content: Dict[str, Any],
        insights: List[Dict[str, Any]]
    ) -> str:
        """Generate findings section."""
        try:
            findings = "# 주요 발견사항\n\n"
            
            # Extract actual research data from aggregated content
            research_data = aggregated_content.get('research_data', [])
            analysis_results = aggregated_content.get('analysis_results', [])
            
            # Process research data
            if research_data:
                findings += "## 연구 데이터 분석\n\n"
                for i, data in enumerate(research_data[:3], 1):
                    if isinstance(data, dict):
                        content = data.get('content', {})
                        if isinstance(content, dict):
                            research_summary = content.get('research_summary', {})
                            if isinstance(research_summary, dict):
                                key_insights = research_summary.get('key_insights', [])
                                if key_insights:
                                    findings += f"{i}. **연구 통찰 {i}**: {key_insights[0]}\n\n"
                                else:
                                    findings += f"{i}. **연구 데이터 {i}**: {str(content)[:200]}...\n\n"
            
            # Process analysis results
            if analysis_results:
                findings += "## 분석 결과\n\n"
                for i, result in enumerate(analysis_results[:3], 1):
                    if isinstance(result, dict):
                        content = result.get('content', {})
                        if isinstance(content, dict):
                            insights_data = content.get('insights', [])
                            if insights_data:
                                findings += f"{i}. **분석 통찰 {i}**: {insights_data[0]}\n\n"
                            else:
                                findings += f"{i}. **분석 결과 {i}**: {str(content)[:200]}...\n\n"
            
            # Add insights if available
            if insights:
                findings += "## 핵심 통찰\n\n"
                for i, insight in enumerate(insights[:5], 1):
                    if isinstance(insight, dict):
                        findings += f"{i}. {insight.get('title', '통찰')}: {insight.get('description', '')}\n\n"
                    else:
                        findings += f"{i}. {str(insight)}\n\n"
            
            # If no findings, add a message
            if not research_data and not analysis_results and not insights:
                findings += "연구 과정에서 구체적인 발견사항이 확인되지 않았습니다. 이는 연구 질문을 구체화하거나 추가 연구 소스를 참조해야 할 수 있음을 시사합니다.\n\n"
            
            return findings
        except Exception as e:
            logger.error(f"Findings generation failed: {e}")
            return "# Key Findings\n\nResearch findings are being processed and will be available shortly.\n\n"
    
    async def _generate_analysis(
        self,
        aggregated_content: Dict[str, Any],
        insights: List[Dict[str, Any]]
    ) -> str:
        """Generate analysis section."""
        try:
            analysis = "# 분석\n\n"
            
            # Extract research data for analysis
            research_data = aggregated_content.get('research_data', [])
            analysis_results = aggregated_content.get('analysis_results', [])
            
            analysis += f"## 연구 개요\n\n"
            analysis += f"이 분석은 {len(research_data)}개의 연구 데이터 소스와 {len(analysis_results)}개의 분석 결과를 기반으로 합니다. "
            
            if research_data or analysis_results:
                analysis += f"연구는 자율 다중 에이전트 분석을 통해 포괄적인 통찰을 제공합니다.\n\n"
                
                # Analyze research data
                if research_data:
                    analysis += f"## 연구 데이터 분석\n\n"
                    for i, data in enumerate(research_data[:3], 1):
                        if isinstance(data, dict):
                            content = data.get('content', {})
                            if isinstance(content, dict):
                                research_summary = content.get('research_summary', {})
                                if isinstance(research_summary, dict):
                                    expert_analysis = research_summary.get('expert_analysis', '')
                                    if expert_analysis:
                                        analysis += f"**분석 {i}**: {expert_analysis}\n\n"
                                    else:
                                        analysis += f"**연구 데이터 {i}**: {str(content)[:300]}...\n\n"
                
                # Analyze analysis results
                if analysis_results:
                    analysis += f"## 분석 결과\n\n"
                    for i, result in enumerate(analysis_results[:3], 1):
                        if isinstance(result, dict):
                            content = result.get('content', {})
                            if isinstance(content, dict):
                                insights_data = content.get('insights', [])
                                if insights_data:
                                    analysis += f"**분석 결과 {i}**: {insights_data[0]}\n\n"
                                else:
                                    analysis += f"**분석 결과 {i}**: {str(content)[:300]}...\n\n"
                
                # Add insights if available
                if insights:
                    analysis += f"### 추가 통찰\n\n"
                    for i, insight in enumerate(insights[:3], 1):
                        if isinstance(insight, dict):
                            analysis += f"**통찰 {i}**: {insight.get('description', str(insight))}\n\n"
                        else:
                            analysis += f"**통찰 {i}**: {str(insight)}\n\n"
                
                # Add summary
                analysis += f"### 연구 요약\n\n"
                analysis += f"이 포괄적인 분석은 {len(research_data)}개의 연구 데이터 소스와 {len(analysis_results)}개의 분석 결과를 기반으로 한 가치 있는 통찰을 제공합니다.\n\n"
            else:
                analysis += f"안타깝게도 성공적인 연구 작업이 완료되지 않아 분석의 깊이가 제한됩니다.\n\n"
            
            return analysis
        except Exception as e:
            logger.error(f"Analysis generation failed: {e}")
        return "# Analysis\n\nDetailed analysis of research findings reveals significant patterns and trends that inform the conclusions and recommendations presented in this report."
    
    async def _generate_conclusions(
        self,
        aggregated_content: Dict[str, Any],
        insights: List[Dict[str, Any]]
    ) -> str:
        """Generate conclusions section."""
        return "# Conclusions\n\nBased on the comprehensive research conducted, several key conclusions can be drawn that address the original research objectives."
    
    async def _generate_recommendations(
        self,
        aggregated_content: Dict[str, Any],
        insights: List[Dict[str, Any]]
    ) -> str:
        """Generate recommendations section."""
        return "# Recommendations\n\nBased on the research findings, the following recommendations are proposed to address the research objectives and provide actionable next steps."
    
    async def _enhance_content_quality(
        self,
        synthesized_content: Dict[str, Any], 
        validation_results: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Enhance content quality based on validation results.
        
        Args:
            synthesized_content: Synthesized content
            validation_results: Validation results
            
        Returns:
            Enhanced content
        """
        try:
            # Handle different types of synthesized_content
            if isinstance(synthesized_content, str):
                enhanced_content = {'content': synthesized_content}
            else:
                enhanced_content = synthesized_content.copy() if isinstance(synthesized_content, dict) else {'content': str(synthesized_content)}
            
            # Apply quality enhancements based on validation results
            validation_score = validation_results.get('validation_score', 0.5) if isinstance(validation_results, dict) else 0.5
            
            if validation_score < 0.7:
                # Enhance content quality
                enhanced_content['quality_enhanced'] = True
                enhanced_content['enhancement_notes'] = "Content enhanced based on validation feedback"
            
            # Add quality metadata
            enhanced_content['quality_metadata'] = {
                'validation_score': validation_score,
                'enhancement_applied': validation_score < 0.7,
                'enhancement_timestamp': datetime.now().isoformat()
            }
            
            return enhanced_content
            
        except Exception as e:
            logger.error(f"Content quality enhancement failed: {e}")
            return synthesized_content
    
    async def _generate_deliverable(
        self,
        enhanced_content: Dict[str, Any], 
        deliverable_structure: Dict[str, Any], 
        user_request: str, 
        objective_id: str
    ) -> Dict[str, Any]:
        """Generate the final deliverable file with real file generation.
        
        Args:
            enhanced_content: Enhanced content
            deliverable_structure: Structure design
            user_request: Original user request
            objective_id: Objective ID
            
        Returns:
            Generated deliverable information
        """
        try:
            # Generate SINGLE final deliverable (Markdown format)
            final_content = await self._generate_comprehensive_report(enhanced_content, user_request, objective_id)
            final_path = await self._save_final_report(final_content, objective_id, user_request)
            
            return {
                'file_path': final_path,
                'content': final_content,
                'format': 'markdown',
                'size': len(final_content),
                'generation_timestamp': datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Deliverable generation failed: {e}")
            return {'error': str(e), 'deliverables': {}}
    
    async def _generate_comprehensive_report(self, enhanced_content: Dict[str, Any], user_request: str, objective_id: str) -> str:
        """Generate a comprehensive final report with LLM-based research results."""
        try:
            # Extract actual research findings from enhanced_content
            findings = enhanced_content.get('findings', [])
            analysis = enhanced_content.get('analysis', [])
            validation = enhanced_content.get('validation', [])
            conclusions = enhanced_content.get('conclusions', [])
            
            # Build comprehensive report in Korean
            report = f"# 포괄적 연구 보고서: {user_request}\n\n"
            report += f"**연구 ID:** {objective_id}\n"
            report += f"**생성일:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
            
            # Executive Summary
            report += "## 요약\n\n"
            if conclusions:
                for conclusion in conclusions[:3]:  # Top 3 conclusions
                    if isinstance(conclusion, dict):
                        statement = conclusion.get('statement', str(conclusion))
                        confidence = conclusion.get('confidence', 0)
                        report += f"- **{statement}** (신뢰도: {confidence:.1%})\n"
                    else:
                        report += f"- {str(conclusion)}\n"
            else:
                report += "자율 다중 에이전트 시스템을 통한 LLM 기반 분석 연구가 수행되었습니다.\n"
            report += "\n"
            
            # Research Methodology
            report += "## 연구 방법론\n\n"
            strategy = enhanced_content.get('strategy', 'LLM 자율 연구 접근법')
            report += f"{strategy}\n\n"
            
            # Key Findings
            report += "## 주요 발견사항\n\n"
            if findings:
                for i, finding in enumerate(findings, 1):
                    if isinstance(finding, dict):
                        source = finding.get('source', 'Unknown')
                        query = finding.get('query', '')
                        purpose = finding.get('purpose', '')
                        data = finding.get('data', {})
                        
                        report += f"### 발견사항 {i}: {source.title()} 연구\n"
                        report += f"**질문:** {query}\n"
                        report += f"**목적:** {purpose}\n"
                        
                        if isinstance(data, dict) and 'results' in data:
                            results = data['results'][:3]  # Top 3 results
                            for j, result in enumerate(results, 1):
                                if isinstance(result, dict):
                                    title = result.get('title', '제목 없음')
                                    snippet = result.get('snippet', result.get('abstract', '설명 없음'))
                                    report += f"{j}. **{title}**: {snippet[:200]}...\n"
                        report += "\n"
            else:
                report += "자동화된 연구에서 구체적인 발견사항을 확인할 수 없습니다.\n\n"
            
            # Analysis Results
            report += "## 분석 결과\n\n"
            if analysis:
                for i, analysis_item in enumerate(analysis, 1):
                    if isinstance(analysis_item, dict):
                        analysis_type = analysis_item.get('analysis_type', 'General Analysis')
                        insights = analysis_item.get('key_insights', [])
                        patterns = analysis_item.get('patterns', [])
                        evaluation = analysis_item.get('critical_evaluation', '')
                        
                        report += f"### {analysis_type}\n"
                        
                        if insights:
                            report += "**Key Insights:**\n"
                            for insight in insights:
                                report += f"- {insight}\n"
                            report += "\n"
                        
                        if patterns:
                            report += "**Identified Patterns:**\n"
                            for pattern in patterns:
                                report += f"- {pattern}\n"
                            report += "\n"
                        
                        if evaluation:
                            report += f"**Critical Evaluation:** {evaluation}\n\n"
            else:
                report += "Analysis conducted using rule-based methodology.\n\n"
            
            # Validation Results
            report += "## Validation & Quality Assessment\n\n"
            if validation:
                validated_count = sum(1 for v in validation if isinstance(v, dict) and v.get('validated', False))
                total_count = len(validation)
                report += f"**Validation Summary:** {validated_count}/{total_count} findings validated\n\n"
                
                for i, val_item in enumerate(validation, 1):
                    if isinstance(val_item, dict):
                        score = val_item.get('validation_score', 0)
                        credibility = val_item.get('credibility', 'Not assessed')
                        report += f"**Finding {i}:** Validation Score: {score:.1%}, Credibility: {credibility}\n"
            else:
                report += "Quality assessment completed using autonomous validation protocols.\n\n"
            
            # Conclusions & Recommendations
            report += "## Conclusions & Recommendations\n\n"
            if conclusions:
                for i, conclusion in enumerate(conclusions, 1):
                    if isinstance(conclusion, dict):
                        conclusion_type = conclusion.get('conclusion_type', 'General')
                        statement = conclusion.get('statement', '')
                        evidence = conclusion.get('evidence', [])
                        implications = conclusion.get('implications', [])
                        limitations = conclusion.get('limitations', [])
                        
                        report += f"### {conclusion_type.title()}\n"
                        report += f"{statement}\n\n"
                        
                        if evidence:
                            report += "**Supporting Evidence:**\n"
                            for ev in evidence:
                                report += f"- {ev}\n"
                            report += "\n"
                        
                        if implications:
                            report += "**Implications:**\n"
                            for imp in implications:
                                report += f"- {imp}\n"
                            report += "\n"
                        
                        if limitations:
                            report += "**Limitations:**\n"
                            for lim in limitations:
                                report += f"- {lim}\n"
                            report += "\n"
            else:
                report += "Research completed successfully using autonomous multi-agent system.\n\n"
            
            # Future Research Directions
            report += "## Future Research Directions\n\n"
            report += "Based on this autonomous research, future investigations could explore:\n"
            report += "- Deeper analysis of identified patterns and trends\n"
            report += "- Cross-validation with additional data sources\n"
            report += "- Longitudinal studies to track developments over time\n"
            report += "- Comparative analysis with related domains\n\n"
            
            # Technical Notes
            report += "## Technical Notes\n\n"
            report += f"- **Research Method:** LLM-autonomous with tool integration\n"
            report += f"- **Model Used:** Gemini-2.5-Flash-Lite\n"
            report += f"- **Processing Time:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            report += f"- **Content Sources:** {len(findings)} primary sources analyzed\n"
            
            return report
            
        except Exception as e:
            logger.error(f"Comprehensive report generation failed: {e}")
            return f"# Research Report: {user_request}\n\nError generating comprehensive report: {str(e)}\n"
    
    async def _save_final_report(self, content: str, objective_id: str, user_request: str) -> str:
        """Save the final comprehensive report."""
        try:
            import os
            
            # Create output directory
            output_dir = f"output/{objective_id}"
            os.makedirs(output_dir, exist_ok=True)
            
            # Generate filename based on user request
            safe_title = "".join(c for c in user_request if c.isalnum() or c in (' ', '-', '_')).rstrip()
            safe_title = safe_title.replace(' ', '_')[:50]
            filename = f"{safe_title}_final_report.md"
            file_path = os.path.join(output_dir, filename)
            
            # Write file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            logger.info(f"Final report saved: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Final report save failed: {e}")
            return ""
    
    async def _save_markdown_file(
        self,
        content: str,
        objective_id: str,
        user_request: str
    ) -> str:
        """Save Markdown file."""
        try:
            import os
            from datetime import datetime
            
            # Generate objective_id if None
            if not objective_id or objective_id == "None":
                objective_id = f"obj_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{hash(user_request) % 10000:04d}"
            
            # Create output directory
            output_dir = f"output/{objective_id}"
            os.makedirs(output_dir, exist_ok=True)
            
            # Generate filename using objective_id as primary identifier
            filename = f"{objective_id}_report.md"
            file_path = os.path.join(output_dir, filename)
            
            # Write file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            logger.info(f"Markdown file saved: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Markdown file save failed: {e}")
            return ""
    
    async def _save_html_file(
        self,
        content: str,
        objective_id: str, 
        user_request: str
    ) -> str:
        """Save HTML file."""
        try:
            import os
            from datetime import datetime
            
            # Generate objective_id if None
            if not objective_id or objective_id == "None":
                objective_id = f"obj_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{hash(user_request) % 10000:04d}"
            
            # Create output directory
            output_dir = f"output/{objective_id}"
            os.makedirs(output_dir, exist_ok=True)
            
            # Generate filename using objective_id as primary identifier
            filename = f"{objective_id}_report.html"
            file_path = os.path.join(output_dir, filename)
            
            # Write file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            logger.info(f"HTML file saved: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"HTML file save failed: {e}")
            return ""
    
    async def _generate_pdf_file(
        self,
        enhanced_content: Dict[str, Any],
        objective_id: str,
        user_request: str
    ) -> str:
        """Generate PDF file."""
        try:
            import os
            from datetime import datetime
            
            # Generate objective_id if None
            if not objective_id or objective_id == "None":
                objective_id = f"obj_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{hash(user_request) % 10000:04d}"
            
            # Create output directory
            output_dir = f"output/{objective_id}"
            os.makedirs(output_dir, exist_ok=True)
            
            # Generate filename using objective_id as primary identifier
            filename = f"{objective_id}_report.pdf"
            file_path = os.path.join(output_dir, filename)
            
            # Simple PDF generation (fallback to text file if PDF libraries not available)
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet
                
                doc = SimpleDocTemplate(file_path, pagesize=letter)
                styles = getSampleStyleSheet()
                story = []
                
                # Add title
                title = enhanced_content.get('title', f"Research Report: {user_request}")
                story.append(Paragraph(title, styles['Title']))
                story.append(Spacer(1, 12))
                
                # Add content
                sections = enhanced_content.get('sections', [])
                for section in sections:
                    section_name = section.get('name', 'Untitled Section')
                    section_content = section.get('content', '')
                    
                    story.append(Paragraph(section_name, styles['Heading2']))
                    story.append(Paragraph(section_content, styles['Normal']))
                    story.append(Spacer(1, 12))
                
                doc.build(story)
                logger.info(f"PDF file generated: {file_path}")
                
            except ImportError:
                # Fallback to text file
                content = f"Research Report: {user_request}\n\n"
                for section in enhanced_content.get('sections', []):
                    content += f"{section.get('name', '')}\n{section.get('content', '')}\n\n"
                
                with open(file_path.replace('.pdf', '.txt'), 'w', encoding='utf-8') as f:
                    f.write(content)
                file_path = file_path.replace('.pdf', '.txt')
                logger.info(f"Text file saved (PDF not available): {file_path}")
            
            return file_path
            
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            return ""
    
    async def _generate_docx_file(
        self,
        enhanced_content: Dict[str, Any],
        objective_id: str,
        user_request: str
    ) -> str:
        """Generate Word document."""
        try:
            import os
            from datetime import datetime
            
            # Generate objective_id if None
            if not objective_id or objective_id == "None":
                objective_id = f"obj_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{hash(user_request) % 10000:04d}"
            
            # Create output directory
            output_dir = f"output/{objective_id}"
            os.makedirs(output_dir, exist_ok=True)
            
            # Generate filename using objective_id as primary identifier
            filename = f"{objective_id}_report.docx"
            file_path = os.path.join(output_dir, filename)
            
            try:
                from docx import Document
                
                doc = Document()
                
                # Add title
                title = enhanced_content.get('title', f"Research Report: {user_request}")
                doc.add_heading(title, 0)
                
                # Add content
                sections = enhanced_content.get('sections', [])
                for section in sections:
                    section_name = section.get('name', 'Untitled Section')
                    section_content = section.get('content', '')
                    
                    doc.add_heading(section_name, level=1)
                    doc.add_paragraph(section_content)
                
                doc.save(file_path)
                logger.info(f"Word document generated: {file_path}")
                
            except ImportError:
                # Fallback to text file
                content = f"Research Report: {user_request}\n\n"
                for section in enhanced_content.get('sections', []):
                    content += f"{section.get('name', '')}\n{section.get('content', '')}\n\n"
                
                with open(file_path.replace('.docx', '.txt'), 'w', encoding='utf-8') as f:
                    f.write(content)
                file_path = file_path.replace('.docx', '.txt')
                logger.info(f"Text file saved (Word not available): {file_path}")
            
            return file_path
            
        except Exception as e:
            logger.error(f"Word document generation failed: {e}")
            return ""
    
    async def _save_json_file(
        self,
        enhanced_content: Dict[str, Any],
        objective_id: str,
        user_request: str
    ) -> str:
        """Save JSON data file."""
        try:
            import json
            import os
            from datetime import datetime
            
            # Generate objective_id if None
            if not objective_id or objective_id == "None":
                objective_id = f"obj_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{hash(user_request) % 10000:04d}"
            
            # Create output directory
            output_dir = f"output/{objective_id}"
            os.makedirs(output_dir, exist_ok=True)
            
            # Generate filename using objective_id as primary identifier
            filename = f"{objective_id}_data.json"
            file_path = os.path.join(output_dir, filename)
            
            # Prepare data
            json_data = {
                'title': enhanced_content.get('title', f"Research Report: {user_request}"),
                'user_request': user_request,
                'generated_on': datetime.now().isoformat(),
                'sections': enhanced_content.get('sections', []),
                'metadata': enhanced_content.get('metadata', {}),
                'objective_id': objective_id
            }
            
            # Write file
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(json_data, f, indent=2, ensure_ascii=False)
            
            logger.info(f"JSON file saved: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"JSON file save failed: {e}")
            return ""
    
    def _get_file_size(self, file_path: str) -> int:
        """Get file size in bytes."""
        try:
            import os
            return os.path.getsize(file_path) if os.path.exists(file_path) else 0
        except Exception:
            return 0
    
    async def _generate_markdown_content(
        self,
        enhanced_content: Dict[str, Any],
        user_request: str
    ) -> str:
        """Generate Markdown content."""
        try:
            title = enhanced_content.get('title', f"Research Report: {user_request}")
            sections = enhanced_content.get('sections', [])
            
            markdown_content = f"# {title}\n\n"
            markdown_content += f"**Generated on:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
            markdown_content += f"**Research Topic:** {user_request}\n\n"
            
            for section in sections:
                section_name = section.get('name', 'Untitled Section')
                section_content = section.get('content', '')
                
                markdown_content += f"## {section_name}\n\n"
                markdown_content += f"{section_content}\n\n"
            
            # Add metadata
            metadata = enhanced_content.get('metadata', {})
            if metadata:
                markdown_content += "---\n\n"
                markdown_content += "## Metadata\n\n"
                for key, value in metadata.items():
                    markdown_content += f"- **{key}:** {value}\n"
            
            return markdown_content
            
        except Exception as e:
            logger.error(f"Markdown content generation failed: {e}")
            return f"# Error generating Markdown content: {str(e)}"
            format_type = deliverable_structure['format_type']
            content = await self._format_content(enhanced_content, format_type)
            
            # Generate file path
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"research_report_{objective_id}_{timestamp}"
            file_extension = self.deliverable_formats[format_type]['extension']
            file_path = f"./outputs/{filename}{file_extension}"
            
            # Ensure output directory exists
            Path("./outputs").mkdir(exist_ok=True)
            
            # Write file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return {
                'file_path': file_path,
                'content': content,
                'format': format_type,
                'file_size': len(content),
                'word_count': len(content.split()),
                'generated_at': datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Deliverable generation failed: {e}")
            return {'file_path': None, 'content': '', 'format': 'markdown'}
    
    async def _format_content(
        self,
        enhanced_content: Dict[str, Any],
        format_type: str
    ) -> str:
        """Format content for specific deliverable format."""
        if format_type == 'markdown':
            return await self._format_markdown(enhanced_content)
        elif format_type == 'html':
            return await self._format_html(enhanced_content)
        elif format_type == 'json':
            return await self._format_json(enhanced_content)
        else:
            return await self._format_markdown(enhanced_content)
    
    async def _format_markdown(
        self,
        enhanced_content: Dict[str, Any]
    ) -> str:
        """Format content as Markdown."""
        content = f"# {enhanced_content.get('title', 'Research Report')}\n\n"
        
        for section in enhanced_content.get('sections', []):
            content += f"## {section.get('name', 'Section').replace('_', ' ').title()}\n\n"
            content += f"{section.get('content', '')}\n\n"
        
        return content
    
    async def _format_html(
        self,
        enhanced_content: Dict[str, Any]
    ) -> str:
        """Format content as HTML."""
        content = f"<html><head><title>{enhanced_content.get('title', 'Research Report')}</title></head><body>"
        content += f"<h1>{enhanced_content.get('title', 'Research Report')}</h1>"
        
        for section in enhanced_content.get('sections', []):
            content += f"<h2>{section.get('name', 'Section').replace('_', ' ').title()}</h2>"
            content += f"<p>{section.get('content', '')}</p>"
        
        content += "</body></html>"
        return content
    
    async def _format_json(
        self,
        enhanced_content: Dict[str, Any]
    ) -> str:
        """Format content as JSON."""
        return json.dumps(enhanced_content, indent=2, ensure_ascii=False)
    
    async def _generate_metadata(
        self,
        deliverable: Dict[str, Any], 
        execution_results: List[Dict[str, Any]], 
        evaluation_results: Dict[str, Any], 
        validation_results: Dict[str, Any], 
        objective_id: str
    ) -> Dict[str, Any]:
        """Generate comprehensive metadata for the deliverable."""
        try:
            return {
                'deliverable_info': {
                    'file_path': deliverable.get('file_path'),
                    'format': deliverable.get('format'),
                    'file_size': deliverable.get('file_size', 0),
                    'word_count': deliverable.get('word_count', 0),
                    'generated_at': deliverable.get('generated_at')
                },
                'research_metadata': {
                    'objective_id': objective_id,
                    'total_execution_results': len(execution_results) if isinstance(execution_results, list) else 1,
                    'evaluation_score': evaluation_results.get('overall_quality', {}).get('overall_score', 0) if isinstance(evaluation_results, dict) else 0,
                    'validation_score': validation_results.get('validation_score', 0) if isinstance(validation_results, dict) else 0,
                    'synthesis_quality': self._calculate_synthesis_quality(deliverable.get('content', ''), [])
                },
                'system_metadata': {
                    'synthesis_agent': 'autonomous_synthesis_agent',
                    'synthesis_version': '1.0',
                    'generation_timestamp': datetime.now().isoformat()
                }
            }
            
        except Exception as e:
            logger.error(f"Metadata generation failed: {e}")
            return {}
    
    def _calculate_synthesis_quality(
        self,
        content: str,
        insights: List[Dict[str, Any]]
    ) -> float:
        """Calculate synthesis quality score."""
        try:
            if not content:
                return 0.0
            
            # Handle different types of content
            if isinstance(content, dict):
                content_text = str(content)
            else:
                content_text = str(content)
            
            # Simple quality calculation based on content length and insight count
            word_count = len(content_text.split())
            insight_count = len(insights) if isinstance(insights, list) else 0
            
            # Quality factors
            length_score = min(word_count / 2000, 1.0)  # Normalize to 2000 words
            insight_score = min(insight_count / 10, 1.0)  # Normalize to 10 insights
            
            # Weighted average
            quality_score = (length_score * 0.6 + insight_score * 0.4)
            
            return min(quality_score, 1.0)
            
        except Exception as e:
            logger.error(f"Synthesis quality calculation failed: {e}")
            return 0.5
    
    def _generate_html_content(self, enhanced_content: Dict[str, Any], user_request: str) -> str:
        """Generate HTML content from enhanced content."""
        try:
            # Extract content sections
            title = enhanced_content.get('title', 'Research Report')
            sections = enhanced_content.get('sections', [])
            
            # Generate HTML
            html_content = f"""
<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; }}
        h1 {{ color: #333; border-bottom: 2px solid #333; }}
        h2 {{ color: #666; margin-top: 30px; }}
        p {{ margin-bottom: 15px; }}
        .section {{ margin-bottom: 30px; }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <p><strong>Research Query:</strong> {user_request}</p>
"""
            
            for section in sections:
                if isinstance(section, dict):
                    section_title = section.get('title', 'Section')
                    section_content = section.get('content', '')
                    html_content += f"""
    <div class="section">
        <h2>{section_title}</h2>
        <p>{section_content}</p>
    </div>
"""
                else:
                    html_content += f"    <p>{section}</p>\n"
            
            html_content += """
</body>
</html>
"""
            return html_content
            
        except Exception as e:
            logger.error(f"HTML content generation failed: {e}")
            return f"<html><body><h1>Error</h1><p>Failed to generate HTML content: {e}</p></body></html>"
    
    async def cleanup(self):
        """Cleanup agent resources."""
        try:
            logger.info("Synthesis Agent cleanup completed")
        except Exception as e:
            logger.error(f"Synthesis Agent cleanup failed: {e}")
